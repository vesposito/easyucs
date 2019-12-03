# coding: utf-8
# !/usr/bin/env python

""" report.py: Easy UCS Deployment Tool """
from easyucs import __author__, __copyright__,  __version__, __status__


import time, datetime, os, sys, gettext, uuid
from docx import Document
from docx.shared import Cm, RGBColor, Mm
from easyucs.report.section import *


class GenericReport:
    def __init__(self, device, inventory, config, language, output_format, page_layout, directory, filename, size):
        self.uuid = uuid.uuid4()
        self.language = language
        self.output_format = output_format
        self.page_layout = page_layout
        self.directory = directory
        self.filename = filename
        self.size = size
        self.timestamp = time.time()

        self.img_path = self.directory + "/" + device.target + "_"
        self.target = device.target

        self.device = device
        self.inventory = inventory
        self.config = config

        self.logger(message="Creating report")

        self.element_list = [] # list of Section, front page, content table, etc.
        self.current_order_id = 0

        self.title = ""
        if self.device.device_type_short == "ucsm":
            self.title = self.config.system[0].name
        elif self.device.device_type_short == "cimc":
            self.title = self.config.admin_networking[0].management_hostname

        if self.output_format == "docx":
            # Open an empty word as a template for some style setting
            # The template layout is letter by default
            self.document = Document('./report/template.docx')
            if self.page_layout.lower() == "a4".lower():
                section = self.document.sections[0]
                section.page_height = Mm(297)
                section.page_width = Mm(210)
                section.left_margin = Mm(25.4)
                section.right_margin = Mm(25.4)
                section.top_margin = Mm(25.4)
                section.bottom_margin = Mm(25.4)
                section.header_distance = Mm(12.7)
                section.footer_distance = Mm(12.7)
        self.install_language()

    def install_language(self):
        # Command to create the pot
        # py C:\Users\mabuelgh\AppData\Local\Programs\Python\Python35\Tools\i18n\pygettext.py -d report .\report.py .\element.py .\subelement.py .\section.py

        if self.language == "en":
            en = gettext.NullTranslations()
            en.install()
        elif self.language == "fr":
            fr = gettext.translation('report', localedir='./locales', languages=['fr'])
            fr.install()

    def get_current_order_id(self):
        self.current_order_id += 1
        return self.current_order_id - 1

    def logger(self, level='info', message="No message"):
        if self.device:
            self.device.logger(level=level, message=message)


class UcsGenericReport(GenericReport):
    def __init__(self, device, inventory, config, language, output_format, page_layout, directory, filename, size):
        GenericReport.__init__(self, device=device, inventory=inventory, config=config, language=language, output_format=output_format, page_layout=page_layout,
                               directory=directory, filename=filename, size=size)

    def recursive_add_in_word_report(self, element):
        # None if add_in_word_report not found
        call = getattr(element, "add_in_word_report", None)
        if call:
            if callable(call):
                if hasattr(element, "content_list"):
                    if element.content_list:
                        element.add_in_word_report()
                elif hasattr(element, "element_list"):
                    if element.element_list:
                        element.add_in_word_report()
                else:
                    element.add_in_word_report()

        if hasattr(element, "content_list"):
            for content in element.content_list:
                self.recursive_add_in_word_report(element=content)

        if hasattr(element, "element_list"):
            for content in element.element_list:
                self.recursive_add_in_word_report(element=content)


class UcsSystemReport(UcsGenericReport):
    def __init__(self, device, inventory, config, language, output_format, page_layout, directory, filename, size):
        UcsGenericReport.__init__(self, device=device, inventory=inventory, config=config, language=language,
                                  output_format=output_format, page_layout=page_layout,
                                  directory=directory, filename=filename, size=size)

        self.element_list.append(
            ReportFrontPage(order_id=self.get_current_order_id(),parent=self,title=self.title,
                            description=_("Created with EasyUCS") + " " + __version__ + " - " +
                                        time.strftime("%d %B %Y"), authors=__author__))

        # We determine the order id of the ReportTableOfContents at the begining (second element to be added but
        # created at the end for obvious reason)
        table_content_order_id = self.get_current_order_id()

        self.element_list.append(
            ClusterOverviewUcsReportSection(order_id=self.get_current_order_id(), parent=self))
        self.element_list.append(
            ArchitectureUcsReportSection(order_id=self.get_current_order_id(), parent=self))
        self.element_list.append(
            EquipmentInventoryUcsReportSection(order_id=self.get_current_order_id(), parent=self))
        self.element_list.append(
            LogicalConfigurationUcsReportSection(order_id=self.get_current_order_id(), parent=self))

        # Always last to append
        self.element_list.append(
            ReportTableOfContents(order_id=table_content_order_id, parent=self, section_list=self.element_list))
        table_of_content = self.element_list[-1]

        # We sort the element list in order to put the table of content at the begining
        self.element_list = sorted(self.element_list, key=lambda element: element.order_id)

        self.logger(level="debug", message="All report elements added to the list")

        if self.output_format == "docx":
            self.logger(message="Generating " + self.output_format + " report")
            self.recursive_add_in_word_report(element=self)

            # Change margins, footer, headers
            section = self.document.sections[0]
            section.right_margin = Cm(2)
            section.left_margin = Cm(2)
            section.different_first_page_header_footer = True
            section.header.paragraphs[0].alignment = 1  # Centered
            section.header.paragraphs[0].text = "Created with EasyUCS " + __version__ + " - " + str(self.title)
            try:
                full_filename = self.directory + '/' + filename + '.docx'
                self.document.save(full_filename)

                # Update the TOC (Windows only)
                # Not used
                # table_of_content.update_toc(full_filename=full_filename)

                self.logger(message="Report generated and saved: " + full_filename)
            except PermissionError:
                self.logger(level="error",
                            message="Report has not been created. Permission denied. Please close Word document.")


class UcsImcReport(UcsGenericReport):
    def __init__(self, device, inventory, config, language, output_format, page_layout, directory, filename, size):
        UcsGenericReport.__init__(self, device=device, inventory=inventory, config=config, language=language,
                                  output_format=output_format, page_layout=page_layout,
                                  directory=directory, filename=filename, size=size)

        self.element_list.append(
            ReportFrontPage(order_id=self.get_current_order_id(), parent=self, title=self.title,
                            description=_("Created with EasyUCS") + " " + __version__ + " - " +
                                        time.strftime("%d %B %Y"), authors=__author__))

        # We determine the order id of the ReportTableOfContents at the begining (second element to be added but
        # created at the end for obvious reason)
        table_content_order_id = self.get_current_order_id()

        self.element_list.append(
            EquipmentInventoryUcsImcReportSection(order_id=self.get_current_order_id(), parent=self))
        # self.element_list.append(
        #     LogicalConfigurationUcsReportSection(order_id=self.get_current_order_id(), parent=self))

        # Always last to append
        self.element_list.append(
            ReportTableOfContents(order_id=table_content_order_id, parent=self, section_list=self.element_list))
        table_of_content = self.element_list[-1]

        # We sort the element list in order to put the table of content at the begining
        self.element_list = sorted(self.element_list, key=lambda element: element.order_id)

        self.logger(level="debug", message="All report elements added to the list")

        if self.output_format == "docx":
            self.logger(message="Generating " + self.output_format + " report")
            self.recursive_add_in_word_report(element=self)

            # Change margins, footer, headers
            section = self.document.sections[0]
            section.right_margin = Cm(2)
            section.left_margin = Cm(2)
            section.different_first_page_header_footer = True
            section.header.paragraphs[0].alignment = 1  # Centered
            section.header.paragraphs[0].text = "Created with EasyUCS " + __version__ + " - " + str(self.title)
            try:
                full_filename = self.directory + '/' + filename + '.docx'
                self.document.save(full_filename)

                # Update the TOC (Windows only)
                # Not used
                # table_of_content.update_toc(full_filename=full_filename)

                self.logger(message="Report generated and saved: " + full_filename)
            except PermissionError:
                self.logger(level="error",
                            message="Report has not been created. Permission denied. Please close Word document.")

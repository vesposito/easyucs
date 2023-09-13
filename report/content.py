# coding: utf-8
# !/usr/bin/env python

""" content.py: Easy UCS Deployment Tool """

import os.path
import sys

from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.opc import constants
from docx.oxml import shared, OxmlElement, ns
from docx.shared import Cm, Pt, RGBColor
from matplotlib import colors

from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, PageBreak, KeepTogether, Table, TableStyle, Spacer
from reportlab.platypus import Image as img
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus.tableofcontents import TableOfContents
from PIL import Image


class GenericReportElement:
    def __init__(self, order_id, parent):
        self.order_id = order_id
        self.parent = parent

        self.report = self.__find_report()

    def __find_report(self):
        current_object = self.parent
        while hasattr(current_object, 'parent') and not hasattr(current_object, 'uuid'):
            current_object = current_object.parent
        if hasattr(current_object, 'uuid'):
            return current_object
        else:
            return None

    def logger(self, level='info', message="No message"):
        if not self.report:
            self.report = self._find_report()

        if self.report:
            self.report.logger(level=level, message=message)


class GenericReportContent(GenericReportElement):
    def __init__(self, order_id, parent, centered=False):
        GenericReportElement.__init__(self, order_id=order_id, parent=parent)
        self.centered = centered  # True or False


class GenericReportHeading(GenericReportContent):
    def __init__(self, order_id, parent, string="", indent=False):
        GenericReportContent.__init__(self, order_id=order_id, parent=parent)
        self.string = string
        if self.indent:
            self.indent = indent
        else:
            self.indent = self.__find_indent()

    def add_in_word_report(self):
        # self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        self.report.document.add_heading(text=self.string, level=self.indent)

    def add_in_json_report(self, content):
        content["Title"] = self.string.replace("\n", "")

    def add_in_pdf_report(self):
        heading_style = ParagraphStyle('heading', alignment=TA_LEFT)
        self.report.pdf_element_list.append(Paragraph(self.string, heading_style))

    def __find_indent(self):
        indent = 1
        current_object = self.parent
        while hasattr(current_object, 'parent') and not hasattr(current_object, 'uuid'):
            current_object = current_object.parent
            indent += 1
        if hasattr(current_object, 'uuid'):
            return indent
        else:
            return None


class GenericReportImage(GenericReportContent):
    def __init__(self, order_id, parent, path, centered=False, size=15, spacing_after=10):
        GenericReportContent.__init__(self, order_id=order_id, parent=parent, centered=centered)
        self.path = path
        self.paragraph = None
        self.size = size
        self.spacing_after = spacing_after

        self.not_found = None

        if not os.path.exists(self.path):
            self.logger(level="warning", message="Image : " + self.path + " not found")
            self.not_found = True

    def add_in_word_report(self):
        # self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        self.paragraph = self.report.document.add_paragraph()
        self.paragraph.alignment = int(bool(self.centered))
        self.paragraph.paragraph_format.space_after = Pt(self.spacing_after)
        if self.not_found:
            self.paragraph.add_run(_("Picture Not Found"))
            self.logger(level="warning", message="Picture " + self.path + " not found: Impossible to add "
                                                                          "it to the report")
        else:
            self.paragraph.add_run().add_picture(self.path, width=Cm(self.size))

    def add_in_json_report(self, content):
        content["ImagePath"] = self.path.replace("\n", "")

    def add_in_pdf_report(self):
        if self.not_found:
            description_style = ParagraphStyle('description', alignment=TA_CENTER)
            self.report.pdf_element_list.append(Paragraph("Picture Not Found", description_style))
        else:
            max_width = 33.972500000000004
            max_height = 38.20583333333334
            img1 = Image.open(self.path)
            width = img1.width * 2.54 / 96
            height = img1.height * 2.54 / 96
            # If the width and height of the image can't fit in a page we find the aspect ratio and downsize the image
            if width > max_width or height > max_height:
                ratio = min(max_width / width, max_height / height)
                width = width * ratio
                height = height * ratio
            image = img(self.path, width/2 * cm, height/2 * cm)
            self.report.pdf_element_list.append(KeepTogether(image))


class GenericReportText(GenericReportContent):
    def __init__(self, order_id, parent, string, centered=False, italicized=False, bolded=False, underlined=False,
                 font="Calibri", font_size=11, color="black", new_paragraph=True, hyper_link=""):
        GenericReportContent.__init__(self, order_id=order_id, parent=parent, centered=centered)
        self.string = string
        self.italicized = italicized
        self.underlined = underlined
        self.bolded = bolded
        self.font = font
        self.font_size = font_size
        self.color = color
        self.new_paragraph = new_paragraph
        self.hyper_link = hyper_link  # The link of the string

    def add_in_word_report(self):
        # self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        if self.new_paragraph:
            paragraph = self.report.document.add_paragraph()
        else:
            paragraph = self.report.document.paragraphs[-1]

        if self.hyper_link:
            # Original: https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
            part = paragraph.part
            r_id = part.relate_to(self.hyper_link, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            # Create the w:hyperlink tag and add needed values
            hyperlink = shared.OxmlElement('w:hyperlink')
            hyperlink.set(shared.qn('r:id'), r_id)

            # Create a w:r element and a new w:rPr element
            new_run = shared.OxmlElement('w:r')
            rPr = shared.OxmlElement('w:rPr')

            # Join all the xml elements together and add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = self.string
            hyperlink.append(new_run)

            # Create a new Run object and add the hyperlink into it
            string = paragraph.add_run()
            string._r.append(hyperlink)

        else:
            string = paragraph.add_run(self.string)

        paragraph.alignment = int(bool(self.centered))
        string.italic = self.italicized
        string.bold = self.bolded
        string.underline = self.underlined
        string.font.size = Pt(self.font_size)
        string.font.name = self.font
        a, b, c = colors.to_rgb(self.color)
        string.font.color.rgb = RGBColor(int(a * 255), int(b * 255), int(c * 255))

    def add_in_json_report(self, content):
        content["Text"] = self.string.replace("\n", "")
        if self.hyper_link:
            content["Hyperlink"] = self.hyper_link

    def add_in_pdf_report(self):
        self.report.pdf_element_list.append(Spacer(1, 1))
        FONT_NAME = "Helvetica"
        if self.bolded:
            FONT_NAME = "Helvetica-Bold"
            self.report.pdf_element_list.append(Spacer(1, 10))
        if self.italicized:
            FONT_NAME = "Helvetica-Oblique"
            self.report.pdf_element_list.append(Spacer(1, 10))

        heading_style = ParagraphStyle("title", alignment=int(bool(self.centered)), fontSize=self.font_size,
                                       textColor=self.color, fontName=FONT_NAME)
        self.report.pdf_element_list.append(Paragraph(self.string, heading_style))


class GenericReportAdmonition(GenericReportContent):
    def __init__(self, order_id, parent, centered=False, font_size=10, level="info", string1="", string2=""):
        """
        :param order_id:
        :param parent:
        :param centered:
        :param font_size:
        :param level:
        :param string1:
        :param string2:
        """

        GenericReportContent.__init__(self, order_id=order_id, parent=parent, centered=centered)
        self.font_size = font_size
        self.level = level
        self.string1 = string1
        self.string2 = string2

        if self.string1:
            if not self.string1.endswith("\n"):
                self.string1 += "\n"

        if level == "ok":
            self.style = "Admonition OK"
        if level == "info":
            self.style = "Admonition Info"
        if level == "warning":
            self.style = "Admonition Warning"
        if level == "error":
            self.style = "Admonition Error"

    def add_in_word_report(self):
        # self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        table = self.report.document.add_table(rows=1, cols=2, style=self.style)

        # Icons: https://www.iconarchive.com/show/small-n-flat-icons-by-paomedia.3.html
        pic = "report/icons/" + self.level + ".png"
        paragraph = table.rows[0].cells[0].paragraphs[0]
        paragraph.alignment = 1
        run = paragraph.add_run()
        run.add_picture(pic, width=Pt(32), height=Pt(32))

        paragraph = table.rows[0].cells[1].paragraphs[0]
        run1 = paragraph.add_run(self.string1)
        run1.bold = True
        if self.string2:
            run2 = paragraph.add_run(self.string2)
            run2.bold = False

        # Set autofit - For more info : https://github.com/python-openxml/python-docx/issues/209
        table.rows[0].cells[0]._tc.tcPr.tcW.type = 'auto'
        table.rows[0].cells[1].width = Cm(13)

        # Change to keep the table on one page 1/2 - Put all the value at False (default values)
        table.style.paragraph_format.keep_together = False
        table.style.paragraph_format.keep_with_next = False

        # Change the font size
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(self.font_size)

        # Change to keep the table on one page 2/2
        # Put the entire table at "keep together" and then put all the rows, except the last one, at "keep with next"
        # Tips found here : https://wordribbon.tips.net/T012975_Keeping_Tables_on_One_Page
        table.style.paragraph_format.keep_together = True
        for row in table.rows[:-1]:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.paragraph_format.keep_with_next = True

    def add_in_json_report(self, content):
        # TODO: Organize the json properly
        content["String1"] = self.string1.replace("\n", "")
        content["String2"] = self.string2.replace("\n", "")
        content["Admonition"] = self.style

    def add_in_pdf_report(self):
        # self.report.pdf_element_list.append(Paragraph("<br/><br/>"))
        self.report.pdf_element_list.append(Spacer(1, 10))
        pic = "report/icons/" + self.level + ".png"
        data = []
        table = []
        image = img(pic, 1 * cm, 1 * cm)
        data.append(image)
        texts = '<font size="9" fontname = "Helvetica-Bold">' + self.string1+'</font><br/>'
        if self.string2:
            texts = texts + '<font size = "9">' + self.string2 + '</font>'
        paragraph = Paragraph(texts)
        data.append(paragraph)
        table.append(data)
        tables = Table(table, colWidths=(1.5*cm, 12*cm))
        table_style = TableStyle([("ALIGN", (0, 0), (1, 1), "LEFT"),
                                  ("VALIGN", (0, 0), (1, 1), "MIDDLE")])
        if self.level == "ok":
            table_style.add("BACKGROUND", (0, 0), (1, 1), "#eaf1dd")
            table_style.add("BOX", (0, 0), (-1, -1), 0.10, "#30cb6b")
        if self.level == "info":
            table_style.add("BACKGROUND", (0, 0), (1, 1), "#dbe5f1")
            table_style.add("BOX", (0, 0), (-1, -1), 0.10, "#4b81c0")
        if self.level == "warning":
            table_style.add("BACKGROUND", (0, 0), (1, 1), "#fce9da")
            table_style.add("BOX", (0, 0), (-1, -1), 0.10, "#edc41d")
        if self.level == "error":
            table_style.add("BACKGROUND", (0, 0), (1, 1), "#f1dcdb")
            table_style.add("BOX", (0, 0), (-1, -1), 0.10, "#de4a39")
        tables.setStyle(table_style)
        self.report.pdf_element_list.append(KeepTogether(tables))


class GenericReportAdmonitionSummary(GenericReportContent):
    def __init__(self, order_id, parent, centered=False, font_size=10, level="info", string_bullets=[]):
        """
        :param order_id:
        :param parent:
        :param centered:
        :param font_size:
        :param level:
        :param string_bullets:
        """

        GenericReportContent.__init__(self, order_id=order_id, parent=parent, centered=centered)
        self.font_size = font_size
        self.level = level
        self.string_bullets = string_bullets

        if level == "ok":
            self.style = "Admonition OK"
        if level == "info":
            self.style = "Admonition Info"
        if level == "warning":
            self.style = "Admonition Warning"
        if level == "error":
            self.style = "Admonition Error"

    def add_in_word_report(self):
        # self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        table = self.report.document.add_table(rows=1, cols=2, style=self.style)

        # Icons: https://www.iconarchive.com/show/small-n-flat-icons-by-paomedia.3.html
        pic = "report/icons/" + self.level + ".png"
        paragraph = table.rows[0].cells[0].paragraphs[0]
        paragraph.alignment = 1
        run = paragraph.add_run()
        run.add_picture(pic, width=Pt(32), height=Pt(32))

        paragraph = table.rows[0].cells[0].add_paragraph()
        paragraph.alignment = 1
        run_text = paragraph.add_run()
        run_text.add_text(self.level.capitalize())
        run_text.bold = True

        paragraph = table.rows[0].cells[1].paragraphs[0]
        run_text = paragraph.add_run()
        if self.level == "error":
            run_text.add_text("Your configuration contains the following " + str(len(self.string_bullets)) +
                              " unsupported critical item(s):")
        elif self.level == "warning":
            run_text.add_text("Your configuration contains the following " + str(self.level) + " item(s):")

        for string_bullet in sorted(self.string_bullets):
            paragraph = table.rows[0].cells[1].add_paragraph()
            paragraph.style = 'List Bullet'
            run1 = paragraph.add_run(string_bullet)
            run1.bold = True

        # Set autofit - For more info : https://github.com/python-openxml/python-docx/issues/209
        table.rows[0].cells[0]._tc.tcPr.tcW.type = 'auto'
        table.rows[0].cells[1].width = Cm(13)

        # Change to keep the table on one page 1/2 - Put all the value at False (default values)
        table.style.paragraph_format.keep_together = False
        table.style.paragraph_format.keep_with_next = False

        # Change the font size
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(self.font_size)

        # Change to keep the table on one page 2/2
        # Put the entire table at "keep together" and then put all the rows, except the last one, at "keep with next"
        # Tips found here : https://wordribbon.tips.net/T012975_Keeping_Tables_on_One_Page
        table.style.paragraph_format.keep_together = True
        for row in table.rows[:-1]:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.paragraph_format.keep_with_next = True

    def add_in_json_report(self, content):
        # TODO: Organize the json properly
        content["StringBullets"] = self.string_bullets
        content["Admonition"] = self.style

    def add_in_pdf_report(self):
        pic = "report/icons/" + self.level + ".png"
        data = []
        table = []
        firstCell = []
        secondCell = []

        # self.report.pdf_element_list.append(Paragraph("<br/><br/>"))
        self.report.pdf_element_list.append(Spacer(1, 25))

        image = img(pic, 1 * cm, 1 * cm)
        admonition_style = ParagraphStyle('description', alignment=TA_CENTER,  fontName='Helvetica-Bold', fontSize=9)
        bullet_style = ParagraphStyle('description', fontName='Helvetica-Bold', fontSize=9, bulletText="•")
        firstCell.append(image)
        firstCell.append(Paragraph(self.level.capitalize(), admonition_style))

        description_style = ParagraphStyle('description', fontSize=9, leading=17)

        if self.level == "error":
            secondCell.append(Paragraph("Your configuration contains the following " + str(len(self.string_bullets)) +
                              " unsupported critical item(s):", description_style))
        elif self.level == "warning":
            secondCell.append(Paragraph("Your configuration contains the following " +
                                        str(self.level) + " item(s):", description_style))
        for string_bullet in sorted(self.string_bullets):
            t = Paragraph(string_bullet, bullet_style)
            secondCell.append(t)

        data.append(firstCell)
        data.append(secondCell)
        table.append(data)

        tables = Table(table, colWidths=(2*cm, 12*cm))
        table_style = TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER"),
                                  ("VALIGN", (0, 0), (1, 1), "MIDDLE")])
        if self.level == "ok":
            table_style.add("BACKGROUND", (0, 0), (1, 1), "#eaf1dd")
            table_style.add("BOX", (0, 0), (-1, -1), 0.10, "#30cb6b")
        if self.level == "info":
            table_style.add("BACKGROUND", (0, 0), (1, 1), "#dbe5f1")
            table_style.add("BOX", (0, 0), (-1, -1), 0.10, "#4b81c0")
        if self.level == "warning":
            table_style.add("BACKGROUND", (0, 0), (1, 1), "#fce9da")
            table_style.add("BOX", (0, 0), (-1, -1), 0.10, "#edc41d")
        if self.level == "error":
            table_style.add("BACKGROUND", (0, 0), (1, 1), "#f1dcdb")
            table_style.add("BOX", (0, 0), (-1, -1), 0.10, "#de4a39")
        tables.setStyle(table_style)
        self.report.pdf_element_list.append(KeepTogether(tables))


class ReportFrontPage(GenericReportElement):
    def __init__(self, order_id, parent, title, description="", authors=""):
        GenericReportElement.__init__(self, order_id=order_id, parent=parent)
        self.authors = authors
        self.description = description
        self.paragraph = None
        self.title = title

    def add_in_word_report(self):
        # self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        page_breaks = 26
        count_break = 0
        # Add line breaks
        for i in range(0, int(page_breaks / 2) - 4):
            self.report.document.add_paragraph()
            count_break += 1

        self.paragraph = self.report.document.add_paragraph()
        self.paragraph.alignment = int(bool(True))
        title = self.paragraph.add_run(self.title)
        title.bold = True
        title.font.size = Pt(28)
        count_break += 1

        self.paragraph = self.report.document.add_paragraph()
        self.paragraph.alignment = int(bool(True))
        description = self.paragraph.add_run(self.description)
        description.italic = True
        self.paragraph = self.report.document.add_paragraph()
        count_break += 2

        # Add line breaks, we adjust to write 2 lines before the end of the page
        for i in range(count_break, page_breaks - 3):
            self.paragraph = self.report.document.add_paragraph()
            count_break += 1
        # alignment = 2 to align right
        self.paragraph.alignment = 2
        author = self.paragraph.add_run(self.authors)

        self.report.document.add_page_break()
        # Add the first header
        self.paragraph = self.report.document.add_paragraph()
        # self.report.document.add_heading(self.report.title, level=0)

    def add_in_json_report(self, content):
        # TODO: Organize the json properly
        content["Title"] = self.title.replace("\n", "")
        content["Description"] = self.description.replace("\n", "")

    def add_in_pdf_report(self):
        page_breaks = 26
        count_break = 0
        styles = getSampleStyleSheet()
        title_list = self.title.split()
        description_list = self.description.split("-")
        for i in range(0, int(page_breaks / 2) - 7):
            self.report.pdf_element_list.append(Paragraph("<br/><br/><br/>", styles['Normal']))
            count_break += 1

        heading_style = ParagraphStyle('title', fontSize=28, alignment=1, leading=24, fontName='Helvetica-Bold')
        self.report.pdf_element_list.append(Paragraph(title_list[0], heading_style))
        if len(title_list) == 4:
            self.report.pdf_element_list.append(Paragraph(title_list[1] + " " + title_list[2] + " " + title_list[3],
                                                          heading_style))

        self.report.pdf_element_list.append(Paragraph("<br/><br/><br/>", styles['Normal']))

        description_style = ParagraphStyle('description', alignment=TA_CENTER, fontName='Helvetica-Oblique')
        self.report.pdf_element_list.append(Paragraph(description_list[0], description_style))
        self.report.pdf_element_list.append(Paragraph(description_list[1], description_style))

        for i in range(count_break, page_breaks - 11):
            self.report.pdf_element_list.append(Paragraph("<br/><br/><br/>", styles['Normal']))

        author_style = ParagraphStyle('author', alignment=TA_RIGHT)
        self.report.pdf_element_list.append(Paragraph(self.authors, author_style))

        self.report.pdf_element_list.append(PageBreak())


class ReportTableOfContents(GenericReportElement):
    def __init__(self, order_id, parent, section_list):
        GenericReportElement.__init__(self, order_id=order_id, parent=parent)
        self.section_list = []
        self.generate_list(list=section_list)
        self.title = _("Table of Contents")

    def add_in_word_report(self):
        # self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        # https://stackoverflow.com/questions/51360649/how-to-update-table-of-contents-in-docx-file-with-python-on-linux
        paragraph = self.report.document.add_paragraph(self.title + ' ')
        paragraph.runs[0].font.size = Pt(18)
        paragraph.runs[0].font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(ns.qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(ns.qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p

    def generate_list(self, list, indent=0):
        # Generate the list of content
        spacing = "    "
        for element in list:
            if hasattr(element, "title"):
                title = element.title
                if indent:
                    for i in range(0, indent):
                        title = spacing + title
                self.section_list.append(title)
            if hasattr(element, "content_list"):
                self.generate_list(list=element.content_list, indent=indent + 1)

    def add_in_pdf_report(self):
        paragraph_style = ParagraphStyle('TOCTitle', fontSize=15, alignment=TA_CENTER, fontName='Helvetica-Bold',
                                         textColor='#6a88c1')
        self.report.pdf_element_list.append(Paragraph("Table of Contents", paragraph_style))
        toc = TableOfContents()
        toc.levelStyles = [
            ParagraphStyle(fontName='Times-Bold', fontSize=13, name='Heading1', leftIndent=20, firstLineIndent=-20,
                           spaceBefore=10, leading=5),
            ParagraphStyle(fontSize=10, name='Heading2', leftIndent=30, firstLineIndent=-20, spaceBefore=5, leading=5),
            ParagraphStyle(fontSize=10, name='Heading3', leftIndent=40, firstLineIndent=-20, spaceBefore=5, leading=5)
        ]
        self.report.pdf_element_list.append(toc)

    def update_toc(self, full_filename):
        if "win" in sys.platform:
            try:
                import win32com.client
            except ModuleNotFoundError:
                self.logger(level="warning", message="You do have a Windows system but you are missing a specific "
                                                     "python library. Please use 'pip install pywin32' on command "
                                                     "line prompt and do this action again. Otherwise, the table of "
                                                     "contents will not be automatically updated")
                return False
            try:
                word = win32com.client.DispatchEx("Word.Application")
                os_path = os.path.sys.path[0]
                doc = word.Documents.Open(os_path + "\\" + full_filename)
                doc.TablesOfContents(1).Update()
                doc.Close(SaveChanges=True)
                word.Quit()
                self.logger(level="debug", message="The table of contents is updated")
            except:
                self.logger(level="error", message="Unable to automatically update the table of contents")
                return False
        else:
            self.logger(level="warning", message="Your system is not Windows, so it's impossible to update the "
                                                 "table of contents automatically. Please go to page 2 of the document "
                                                 "and right-click on the text and click 'update field'")
            return False

# coding: utf-8
# !/usr/bin/env python

""" content.py: Easy UCS Deployment Tool """
from easyucs import __author__, __copyright__,  __version__, __status__


from docx.shared import Cm, Pt, RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from matplotlib import colors
import os.path
import sys
from docx.oxml import shared, OxmlElement, ns
from docx.opc import constants


class GenericReportElement():
    def __init__(self, order_id, parent):
        self.order_id = order_id
        self.parent = parent

        self.report = self.__find_report()

    def __find_report(self):
        current_object = self.parent
        while hasattr(current_object, 'parent') and not hasattr(current_object, 'timestamp'):
            current_object = current_object.parent
        if hasattr(current_object, 'timestamp'):
            return current_object
        else:
            return None

    def logger(self, level='info', message="No message"):
        if not self.report:
            self.report = self._find_report()

        if self.report:
            self.report.logger(level=level, message=message)


class GenericReportContent(GenericReportElement):
    def __init__(self, order_id, parent, centered=False):
        GenericReportElement.__init__(self, order_id=order_id, parent=parent)
        self.centered = centered  # True or False


class GenericReportHeading(GenericReportContent):
    def __init__(self, order_id, parent, string, indent=False):
        GenericReportContent.__init__(self, order_id=order_id, parent=parent)
        self.string = string
        if self.indent:
            self.indent = indent
        else:
            self.indent = self.__find_indent()

    def add_in_word_report(self):
        #self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        self.report.document.add_heading(text=self.string, level=self.indent)

    def __find_indent(self):
        indent = 1
        current_object = self.parent
        while hasattr(current_object, 'parent') and not hasattr(current_object, 'timestamp'):
            current_object = current_object.parent
            indent += 1
        if hasattr(current_object, 'timestamp'):
            return indent
        else:
            return None


class GenericReportText(GenericReportContent):
    def __init__(self, order_id, parent, string, centered=False, italicized=False, bolded=False, underlined=False,
                 font="Calibri", font_size=11, color="black", new_paragraph=True, hyper_link=""):
        GenericReportContent.__init__(self, order_id=order_id, parent=parent, centered=centered)
        self.string = string
        self.italicized = italicized
        self.underlined = underlined
        self.bolded = bolded
        self.font = font
        self.font_size = font_size
        self.color = color
        self.new_paragraph = new_paragraph
        self.hyper_link = hyper_link  # The link of the string

    def add_in_word_report(self):
        #self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        if self.new_paragraph:
            paragraph = self.report.document.add_paragraph()
        else:
            paragraph = self.report.document.paragraphs[-1]

        if self.hyper_link:
            # Original: https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
            part = paragraph.part
            r_id = part.relate_to(self.hyper_link, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            # Create the w:hyperlink tag and add needed values
            hyperlink = shared.OxmlElement('w:hyperlink')
            hyperlink.set(shared.qn('r:id'), r_id)

            # Create a w:r element and a new w:rPr element
            new_run = shared.OxmlElement('w:r')
            rPr = shared.OxmlElement('w:rPr')

            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = self.string
            hyperlink.append(new_run)

            # Create a new Run object and add the hyperlink into it
            string = paragraph.add_run()
            string._r.append(hyperlink)

        else:
            string = paragraph.add_run(self.string)

        paragraph.alignment = int(bool(self.centered))
        string.italic = self.italicized
        string.bold = self.bolded
        string.underline = self.underlined
        string.font.size = Pt(self.font_size)
        string.font.name = self.font
        a, b, c = colors.to_rgb(self.color)
        string.font.color.rgb = RGBColor(int(a * 255), int(b * 255), int(c * 255))


class GenericReportImage(GenericReportContent):
    def __init__(self, order_id, parent, path, centered=False, size=15, spacing_after=10):
        GenericReportContent.__init__(self, order_id=order_id, parent=parent, centered=centered)
        self.path = path
        self.size = size
        self.spacing_after = spacing_after

        self.not_found = None

        if not os.path.exists(self.path):
            self.logger(level="warning", message="Image : " + self.path + " not found")
            self.not_found = True

    def add_in_word_report(self):
        #self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        self.paragraph = self.report.document.add_paragraph()
        self.paragraph.alignment = int(bool(self.centered))
        self.paragraph.paragraph_format.space_after = Pt(self.spacing_after)
        if self.not_found:
            self.paragraph.add_run(_("Picture Not Found"))
            self.logger(level="warning", message="Picture " + self.path + " not found: Impossible to add "
                                                                          "it to the report")
        else:
            self.paragraph.add_run().add_picture(self.path, width=Cm(self.size))


class ReportTableOfContents(GenericReportElement):
    def __init__(self, order_id, parent, section_list):
        GenericReportElement.__init__(self, order_id=order_id, parent=parent)
        self.section_list = []
        self.generate_list(list=section_list)
        self.title = _("Table of Contents")

    def add_in_word_report(self):
        #self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        # https://stackoverflow.com/questions/51360649/how-to-update-table-of-contents-in-docx-file-with-python-on-linux
        paragraph = self.report.document.add_paragraph(self.title + ' ')
        paragraph.runs[0].font.size = Pt(18)
        paragraph.runs[0].font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(ns.qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(ns.qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p

    def generate_list(self, list, indent=0):
        # Generate the list of content
        spacing = "    "
        for element in list:
            if hasattr(element, "title"):
                title = element.title
                if indent:
                    for i in range(0, indent):
                        title = spacing + title
                self.section_list.append(title)
            if hasattr(element, "content_list"):
                self.generate_list(list=element.content_list, indent=indent + 1)

    def update_toc(self, full_filename):
        if "win" in sys.platform:
            try:
                import win32com.client
            except ModuleNotFoundError:
                self.logger(level="warning", message="You do have a Windows system but you are missing a specific "
                                                     "python library. Please use 'pip install pywin32' on command "
                                                     "line prompt and do this action again. Otherwise, the table of "
                                                     "contents will not be automatically updated")
                return False
            try:
                word = win32com.client.DispatchEx("Word.Application")
                os_path = os.path.sys.path[0]
                doc = word.Documents.Open(os_path + "\\" + full_filename)
                doc.TablesOfContents(1).Update()
                doc.Close(SaveChanges=True)
                word.Quit()
                self.logger(level="debug", message="The table of contents is updated")
            except:
                self.logger(level="error", message="Unable to automatically update the table of contents")
                return False
        else:
            self.logger(level="warning", message="Your system is not Windows, so it's impossible to update the "
                                                 "table of contents automatically. Please go to page 2 of the document "
                                                 "and right-click on the text and click 'update field'")
            return False


class ReportFrontPage(GenericReportElement):
    def __init__(self, order_id, parent, title, description="", authors=""):
        GenericReportElement.__init__(self, order_id=order_id, parent=parent)
        self.title = title
        self.description = description
        self.authors = authors

    def add_in_word_report(self):
        #self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        page_breaks = 26
        count_break = 0
        # Add line breaks
        for i in range(0, int(page_breaks/2) - 4):
            self.report.document.add_paragraph()
            count_break += 1

        self.paragraph = self.report.document.add_paragraph()
        self.paragraph.alignment = int(bool(True))
        title = self.paragraph.add_run((self.title))
        title.bold = True
        title.font.size = Pt(28)
        count_break += 1

        self.paragraph = self.report.document.add_paragraph()
        self.paragraph.alignment = int(bool(True))
        description = self.paragraph.add_run((self.description))
        description.italic = True
        self.paragraph = self.report.document.add_paragraph()
        count_break += 2

        # Add line breaks, we adjust to write 2 lines before the end of the page
        for i in range(count_break, page_breaks -3):
            self.paragraph = self.report.document.add_paragraph()
            count_break += 1
        # alignment = 2 to align right
        self.paragraph.alignment = 2
        author = self.paragraph.add_run((self.authors))

        self.report.document.add_page_break()
        # Add the first header
        self.paragraph = self.report.document.add_paragraph()
        #self.report.document.add_heading(self.report.title, level=0)

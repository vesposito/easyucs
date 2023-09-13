# coding: utf-8
# !/usr/bin/env python

""" report.py: Easy UCS Deployment Tool """

import gettext
import uuid

from docx import Document
from docx.shared import Cm, Mm

import export
from __init__ import __version__
from repository.metadata import ReportMetadata

from reportlab.platypus import SimpleDocTemplate, PageTemplate, Paragraph
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus.frames import Frame
from reportlab.lib.units import cm


class MyDocTemplate(SimpleDocTemplate):
    def __init__(self, filename, **kw):
        self.allowSplitting = 0
        SimpleDocTemplate.__init__(self, filename, **kw)
        template = PageTemplate('normal', [Frame(2.5 * cm, 2.5 * cm, 15 * cm, 25 * cm, id='F1')])
        self.addPageTemplates(template)

    def afterFlowable(self, flowable):
        # "Registers TOC entries."
        if flowable.__class__.__name__ == 'Paragraph':
            text = flowable.getPlainText()
            style = flowable.style.name
            if style == 'Heading1':
                level = 0
                key = text
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(key, key, 0)
            elif style == 'Heading2':
                level = 1
                key = text
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(key, key, 1)
            elif style == 'Heading3':
                level = 2
                key = text
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(key, key, 2)
            # H4 and H5 are for bookmarks, so we don't assign level for them
            elif style == 'Heading4':
                key = text
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(key, key, 3)
                return
            elif style == 'Heading5':
                key = text
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(key, key, 4)
                return
            elif style == 'Heading6':
                key = text
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(key, key, 5)
                return
            else:
                return
            E = [level, text, self.page]
            # if we have a bookmark name append that to our notify data
            bn = getattr(flowable, '_bookmarkName', None)
            if bn is not None:
                E.append(bn)
            self.notify('TOCEntry', tuple(E))


class GenericReport:
    def __init__(self, parent=None, device=None, inventory=None, config=None, language="en", output_format="docx",
                 page_layout="letter", directory=".", size="full", report_type=None):
        self.config = config
        self.device = device
        self.directory = directory
        self.inventory = inventory
        self.parent = parent
        self.size = size
        self.uuid = uuid.uuid4()

        # Needs to be created after UUID
        kwargs = {
            "parent": self,
            "language": language,
            "output_format": output_format,
            "page_layout": page_layout,
            "size": size
        }
        if config:
            kwargs["config_uuid"] = config.uuid
        if inventory:
            kwargs["inventory_uuid"] = inventory.uuid
        if report_type:
            kwargs["report_type"] = report_type
        self.metadata = ReportMetadata(**kwargs)

        self.target = device.target
        self.img_path = self.directory + "/" + device.target + "_"

        self.document = None
        self.element_list = []  # list of section, front page, content table, etc.
        self.current_order_id = 0
        self.toc_order_id = 0
        self.toc = None
        if not hasattr(self, "title"):
            self.title = ""

        self.install_language()

        if self.metadata.output_format == "docx":
            # Open an empty Word as a template for some style setting
            # The template layout is letter by default
            template_ref = open("./report/template.docx", "rb")
            self.document = Document(template_ref)
            template_ref.close()

            section = self.document.sections[0]
            if self.metadata.page_layout.lower() == "a4":
                section.page_height = Mm(297)
                section.page_width = Mm(210)
                section.left_margin = Cm(2)  # Default A4 margin is Mm(25.4)
                section.right_margin = Cm(2)  # Default A4 margin is Mm(25.4)
                section.top_margin = Mm(25.4)
                section.bottom_margin = Mm(25.4)
                section.header_distance = Mm(12.7)
                section.footer_distance = Mm(12.7)

            # Setting header of report
            section.different_first_page_header_footer = True
            section.header.paragraphs[0].alignment = 1  # Centered
            section.header.paragraphs[0].text = _("Created with EasyUCS") + " " + __version__ + " - " + str(self.title)

        elif self.metadata.output_format == "json":
            self.document = {}
            header_json = {
                "metadata": [export.generate_json_metadata_header(file_type="report", report=self)]
            }
            self.document["easyucs"] = header_json
            self.document["report"] = {}

        elif self.metadata.output_format == "pdf":
            self.pdf_element_list = []
            self.document = MyDocTemplate

        else:
            self.logger(level="error", message=f"Invalid output format. '{self.metadata.output_format}' format not "
                                               "supported")

    def install_language(self):
        if self.metadata.language == "en":
            en = gettext.NullTranslations()
            en.install()
        elif self.metadata.language == "fr":
            fr = gettext.translation('report', localedir='./locales', languages=['fr'])
            fr.install()

    def get_current_order_id(self):
        self.current_order_id += 1
        return self.current_order_id - 1

    def logger(self, level='info', message="No message"):
        if self.device:
            self.device.logger(level=level, message=message)

    def recursive_add_in_word_report(self, element):
        # None if add_in_word_report not found
        call = getattr(element, "add_in_word_report", None)
        if call:
            if callable(call):
                if hasattr(element, "content_list"):
                    if element.content_list:
                        element.add_in_word_report()
                elif hasattr(element, "element_list"):
                    if element.element_list:
                        element.add_in_word_report()
                else:
                    element.add_in_word_report()

        if hasattr(element, "content_list"):
            for content in element.content_list:
                self.recursive_add_in_word_report(element=content)

        if hasattr(element, "element_list"):
            for content in element.element_list:
                self.recursive_add_in_word_report(element=content)

    def recursive_add_in_json_report(self, element, document_json):
        # None if add_in_json_report not found
        call = getattr(element, "add_in_json_report", None)
        if call:
            if callable(call):
                if hasattr(element, "content_list"):
                    if element.content_list:
                        element.add_in_json_report(document_json)
                elif hasattr(element, "element_list"):
                    if element.element_list:
                        element.add_in_json_report(document_json)
                else:
                    element.add_in_json_report(document_json)
        if hasattr(element, "content_list"):
            for content in element.content_list:
                document_json[type(content).__name__ + str(content.order_id)] = {}
                self.recursive_add_in_json_report(element=content,
                                                  document_json=document_json[
                                                      type(content).__name__ + str(content.order_id)])

        if hasattr(element, "element_list"):
            for content in element.element_list:
                document_json[type(content).__name__ + str(content.order_id)] = {}
                self.recursive_add_in_json_report(element=content,
                                                  document_json=document_json[
                                                      type(content).__name__ + str(content.order_id)])

    def recursive_add_in_pdf_report(self, element):
        # None if add_in_pdf_report not found
        call = getattr(element, "add_in_pdf_report", None)
        if call:
            if callable(call):
                if hasattr(element, "content_list"):
                    if element.content_list:
                        element.add_in_pdf_report()
                elif hasattr(element, "element_list"):
                    if element.element_list:
                        element.add_in_pdf_report()
                else:
                    element.add_in_pdf_report()

        if hasattr(element, "content_list"):
            for content in element.content_list:
                self.recursive_add_in_pdf_report(element=content)

        if hasattr(element, "element_list"):
            for content in element.element_list:
                self.recursive_add_in_pdf_report(element=content)

    def _reset_pdf_sequence(self):
        """
        Resets the sequences used for generating PDF reports.
        This function appends sequence reset tags to the PDF element list, allowing the sequences in the PDF
        report to start from 1 each time a new pdf report is created.
        """
        # Implemented to prevent EASYUCS-911
        if hasattr(self, "pdf_element_list"):
            self.pdf_element_list.append(Paragraph("<seqreset id='h1'/>", ParagraphStyle('body')))
            self.pdf_element_list.append(Paragraph("<seqreset id='h2'/>", ParagraphStyle('body')))
            self.pdf_element_list.append(Paragraph("<seqreset id='h3'/>", ParagraphStyle('body')))
            self.pdf_element_list.append(Paragraph("<seqreset id='h4'/>", ParagraphStyle('body')))
            self.pdf_element_list.append(Paragraph("<seqreset id='h5'/>", ParagraphStyle('body')))
            self.pdf_element_list.append(Paragraph("<seqreset id='h6'/>", ParagraphStyle('body')))

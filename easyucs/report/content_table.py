# coding: utf-8
# !/usr/bin/env python

""" content_table.py: Easy UCS Deployment Tool """
from easyucs import __author__, __copyright__,  __version__, __status__


from docx.shared import Cm, Pt
from easyucs.report.content import GenericReportContent
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from ipaddress import IPv4Address, IPv6Address

import os.path

from easyucs.draw.object import GenericDrawObject


class GenericReportTable(GenericReportContent):
    def __init__(self, order_id, parent, row_number, column_number, centered=False, cells_list=[],
                 style="Light Grid Accent 1", autofit=True, font_size=10):
        """
        :param order_id:
        :param parent:
        :param row_number:
        :param column_number:
        :param centered:
        :param cells_list: example : list = [[1,2,3],[2,3,4]]
        :param style:
        :param autofit:
        :param font_size:
        """

        GenericReportContent.__init__(self, order_id=order_id, parent=parent, centered=centered)
        self.row_number = row_number
        self.column_number = column_number
        self.cells_list = cells_list
        self.style = style
        self.autofit = autofit
        self.font_size = font_size

        if len(self.cells_list) == self.row_number:
            self.cells_type = "row"
        elif len(self.cells_list) == self.column_number:
            self.cells_type = "column"
            self.__format_column_to_row()
        else:
            self.logger(level="warning",
                        message="Cells in the table are not the same size as the number of column or row")
            return False

        self.clean_empty()

    def clean_empty(self):
        # Column
        columns_to_delete = []
        for column_index in range(len(self.cells_list[0])):
            empty = True
            for row in self.cells_list:
                if not empty:
                    continue
                if row == self.cells_list[0]:
                    # We skip the first row because it can't be empty in a "row" cells_type table
                    continue
                if row[column_index] or row[column_index] == 0:
                    empty = False
            if empty:
                columns_to_delete.append(column_index)

        # Ajust list of index column to delete because we delete column one by one. So if I need to delete the
        # 2 and 3 column, once the 2nd column is deleted the 3rd column is now the 2nd
        for i in range(len(columns_to_delete)):
            columns_to_delete[i] = columns_to_delete[i] - i

        for column_to_delete in columns_to_delete:
            for row in self.cells_list:
                row.pop(column_to_delete)

        self.row_number = len(self.cells_list)
        self.column_number = len(self.cells_list[0])

        # Rows
        rows_to_delete = []
        for row_index in range(len(self.cells_list)):
            if row_index != 0:
                empty = True
                for column_in_row in self.cells_list[row_index]:
                    if isinstance(column_in_row, type(None)):
                        # Avoid issue with None value that is misinterpret by Word. We change None to empty string
                        self.cells_list[row_index][self.cells_list[row_index].index(column_in_row)] = ""
                    if not empty:
                        continue
                    if column_in_row == self.cells_list[row_index][0]:
                        continue
                    if column_in_row or column_in_row == 0:
                        empty = False
                if empty:
                    rows_to_delete.append(row_index)
        # Ajust list of index row to delete because we delete row one by one. So if I need to delete the
        # 2 and 3 row, once the 2nd row is deleted the 3rd row is now the 2nd
        for i in range(len(rows_to_delete)):
            rows_to_delete[i] = rows_to_delete[i] - i

        for row_to_delete in rows_to_delete:
            self.cells_list.pop(row_to_delete)

        self.row_number = len(self.cells_list)
        self.column_number = len(self.cells_list[0])

    def add_in_word_report(self):
        # self.logger(level="debug", message="Adding " + self.__class__.__name__ + " in word")
        table = self.report.document.add_table(rows=self.row_number, cols=self.column_number,
                                               style=self.style)
        table.alignment = int(bool(self.centered))

        for i in range(len(self.cells_list)):
            row = self.cells_list[i]
            row = list(filter(None.__ne__, row))  # Remove None values
            row_cells = table.rows[i].cells
            for j in range(0, len(row)):
                row_cells[j].text = str(row[j])
                row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if self.autofit:
                    # For more info : https://github.com/python-openxml/python-docx/issues/209
                    row_cells[j]._tc.tcPr.tcW.type = 'auto'

        # Change to keep the table on one page 1/2 - Put all the value at False (default values)
        table.style.paragraph_format.keep_together = False
        table.style.paragraph_format.keep_with_next = False

        # Change the font size
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(self.font_size)

        # Change to keep the table on one page 2/2
        # Put the entire table at "keep together" and then put all the rows, except the last one, at "keep with next"
        # Tips found here : https://wordribbon.tips.net/T012975_Keeping_Tables_on_One_Page
        table.style.paragraph_format.keep_together = True
        for row in table.rows[:-1]:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.paragraph_format.keep_with_next = True

    def __format_column_to_row(self):
        column_cells_list = self.cells_list
        row_cells_list = []
        for i in range(0, self.row_number):
            row_cells_list.append([])

        for i in range(0, len(column_cells_list)):
            for j in range(0, self.row_number):
                row_cells_list[j].append(column_cells_list[i][j])

        self.cells_list = row_cells_list
        self.cells_type = "row"

    @staticmethod
    def get_name_and_sku(inventory_object):
        """
        Generate proper name depending on available attributes
        :param inventory_object:
        :return: key
        """
        key = ""

        if hasattr(inventory_object, "name"):
            if inventory_object.name:
                key = inventory_object.name
                if inventory_object.sku:
                    key = inventory_object.name + ' (' + inventory_object.sku + ')'
                elif inventory_object.model and inventory_object.model != inventory_object.name:
                    key = inventory_object.name + ' (' + inventory_object.model + ')'
        if not key:
            if inventory_object.sku:
                key = inventory_object.sku
            elif inventory_object.model:
                key = inventory_object.model

        return key


class FiUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, fi, centered=False):
        rows = [[_("Description"),_("Value")]]
        rows.append([_("Fabric"), fi.id])
        rows.append([_("SKU"), fi.sku])
        rows.append([_("Model"), fi.name])
        rows.append([_("Serial Number"), fi.serial])
        rows.append([_("Firmware"), fi.firmware_package_version])


        # We have to get the JSON file of the device to get informations about the ports
        # We use a DrawingObject because they already have a method to get the JSON file efficiently
        draw = GenericDrawObject(parent=fi)
        draw._get_json_file()
        rear_ports = None
        if draw.json_file:
            if "rear_ports" in draw.json_file:
                rear_ports = draw.json_file['rear_ports']
        if rear_ports:
            port_dict = {}
            for port in rear_ports.items():
                if port[1]['port_type'] in port_dict:
                    port_dict[port[1]['port_type']] += 1
                else:
                    port_dict.update({port[1]['port_type']: 1})
            port_info = "\n".join([(str(i[1]) + "x " + str(i[0]).upper()) for i in port_dict.items()])
            rows.append([_("Ports"), port_info])

        if fi.power_supplies:
            psu_dict = {}
            for psu in fi.power_supplies:
                key = self.get_name_and_sku(psu)
                if key in psu_dict:
                    psu_dict[key] += 1
                else:
                    psu_dict.update({key: 1})
            psu_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in psu_dict.items()])
            rows.append([_("Power Supplies"), psu_models])
            
        if fi.expansion_modules:
            expansion_dict = {}
            for expansion in fi.expansion_modules:
                key = self.get_name_and_sku(expansion)
                if key in expansion_dict:
                    expansion_dict[key] += 1
                else:
                    expansion_dict.update({key: 1})
            gem_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in expansion_dict.items()])
            rows.append([_("Expansion Modules"), gem_models])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class FexUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, fex, centered=False):
        rows = [[_("Description"),_("Value")]]
        rows.append([_("FEX"), fex.id])
        rows.append([_("Fabric"), fex.switch_id])
        rows.append([_("SKU"), fex.sku])
        rows.append([_("Model"), fex.name])
        rows.append([_("Serial Number"), fex.serial])

        # We have to get the JSON file of the device to get informations about the ports
        # We use a DrawingObject because they already have a method to get the JSON file efficiently
        draw = GenericDrawObject(parent=fex)
        draw._get_json_file()
        rear_ports = draw.json_file['rear_ports']
        if rear_ports:
            port_dict = {}
            for port in rear_ports.items():
                if port[1]['port_type'] in port_dict:
                    port_dict[port[1]['port_type']] += 1
                else:
                    port_dict.update({port[1]['port_type']: 1})
            port_info = "\n".join([(str(i[1]) + "x " + str(i[0]).upper()) for i in port_dict.items()])
            rows.append([_("Ports"), port_info])

        if fex.power_supplies:
            psu_dict = {}
            for psu in fex.power_supplies:
                key = self.get_name_and_sku(psu)
                if key in psu_dict:
                    psu_dict[key] += 1
                else:
                    psu_dict.update({key: 1})
            psu_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in psu_dict.items()])
            rows.append([_("Power Supplies"), psu_models])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class RackUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, rack, centered=False):
        rows = [[_("Description"),_("Value")]]
        rows.append([_("Rack ID"), rack.id])
        rows.append([_("SKU"), rack.sku])
        rows.append([_("Model"), rack.name])
        rows.append([_("Serial Number"), rack.serial])

        memory_info = ""
        if rack.memory_arrays:
            memory_dict = {}
            for array in rack.memory_arrays:
                for unit in array.memory_units:
                    if not unit.capacity:
                        continue
                    if unit.sku:
                        if unit.clock and unit.type:
                            key = str(int(unit.capacity / 1024)) + 'GB ' + str(unit.type) + ' ' + str(unit.clock) + 'MHz (' + unit.sku + ')'
                        elif unit.clock:
                            key = str(int(unit.capacity / 1024)) + 'GB ' + str(unit.clock) + 'MHz (' + unit.sku + ')'
                        else:
                            key = str(int(unit.capacity / 1024)) + 'GB ' + str(unit.type) + ' (' + unit.sku + ')'
                    else:
                        if unit.clock and unit.type:
                            key = str(int(unit.capacity / 1024)) + 'GB ' + str(unit.type) + ' ' + str(unit.clock) + 'MHz'
                        elif unit.clock:
                            key = str(int(unit.capacity / 1024)) + 'GB ' + str(unit.clock) + 'MHz'
                        else:
                            key = str(int(unit.capacity / 1024)) + 'GB' + str(unit.type)
                    if key in memory_dict:
                        memory_dict[key] += 1
                    else:
                        memory_dict.update({key: 1})
            memory_info = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in memory_dict.items()])
        rows.append([_("Memory"), str(rack.memory_total_marketing) + "\n" + memory_info])

        if rack.adaptors:
            adaptor_dict = {}
            for adaptor in rack.adaptors:
                key = self.get_name_and_sku(adaptor)
                if key in adaptor_dict:
                    adaptor_dict[key] += 1
                else:
                    adaptor_dict.update({key: 1})
            adaptor_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in adaptor_dict.items()])
            rows.append([_("Adaptors"), adaptor_models])

        cpu_model = ""
        cores = 0
        if rack.cpus:
            cpu_dict = {}

            for cpu in rack.cpus:
                key = self.get_name_and_sku(cpu)
                if key in cpu_dict:
                    cpu_dict[key] += 1
                else:
                    cpu_dict.update({key: 1})
                if cpu.cores:
                    cores += int(cpu.cores)

            cpu_model = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in cpu_dict.items()])

        rows.append([_("CPUs"), cpu_model])
        rows.append([_("Cores"), cores])

        gpu_models = ""
        if rack.gpus:
            gpu_dict = {}
            for gpu in rack.gpus:
                key = self.get_name_and_sku(gpu)
                if key in gpu_dict:
                    gpu_dict[key] += 1
                else:
                    gpu_dict.update({key: 1})
            gpu_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in gpu_dict.items()])
            rows.append([_("GPUs"), gpu_models])

        storage = ""
        drives = ""
        if rack.storage_controllers:
            storage_dict = {}
            drives_dict = {}
            for controller in rack.storage_controllers:
                key = self.get_name_and_sku(controller)
                if key in storage_dict:
                    storage_dict[key] += 1
                else:
                    storage_dict.update({key: 1})

                for drive in controller.disks:
                    key_drive = self.get_name_and_sku(drive)
                    if key_drive in drives_dict:
                        drives_dict[key_drive] += 1
                    else:
                        drives_dict.update({key_drive: 1})

            for drive in rack.nvme_drives:
                key_drive = self.get_name_and_sku(drive)
                if key_drive in drives_dict:
                    drives_dict[key_drive] += 1
                else:
                    drives_dict.update({key_drive: 1})

            storage = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in storage_dict.items()])
            drives = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in drives_dict.items()])
        rows.append([_("Storage Controllers"), storage])
        rows.append([_("Drives"), drives])

        flash_cards = ""
        if rack.storage_flexflash_controllers:
            flash_dict = {}
            for controller in rack.storage_flexflash_controllers:
                for card in controller.flexflash_cards:
                    if card.capacity_marketing in flash_dict:
                        flash_dict[card.capacity_marketing] += 1
                    else:
                        flash_dict.update({card.capacity_marketing: 1})
            flash_cards = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in flash_dict.items()])
        rows.append([_("FlexFlash SD Cards"), flash_cards])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class ChassisUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, chassis, centered=False):
        rows = [[_("Description"),_("Value")]]
        rows.append([_("Chassis ID"), chassis.id])
        rows.append([_("SKU"), chassis.sku])
        rows.append([_("Model"), chassis.name])
        rows.append([_("Serial Number"), chassis.serial])
        rows.append([_("Slots Used/Total"), str(chassis.slots_populated) + "/" + str(chassis.slots_max)])
        rows.append([_("Slots Free (full-size)"), chassis.slots_free_full])
        rows.append([_("Slots Free (half-size)"), chassis.slots_free_half])


        if chassis.io_modules:
            io_dict = {}
            for io_module in chassis.io_modules:
                key = self.get_name_and_sku(io_module)
                if key in io_dict:
                    io_dict[key] += 1
                else:
                    io_dict.update({key: 1})
            io_modules_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in io_dict.items()])
            rows.append([_("IO Modules"), io_modules_models])
            
        if chassis.system_io_controllers:
            sioc_dict = {}
            for sioc_controller in chassis.system_io_controllers:
                key = self.get_name_and_sku(sioc_controller)
                if key in sioc_dict:
                    sioc_dict[key] += 1
                else:
                    sioc_dict.update({key: 1})
            sioc_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in sioc_dict.items()])
            rows.append([_("SIOCs"), sioc_models])
            
        if chassis.power_supplies:
            psu_dict = {}
            for psu in chassis.power_supplies:
                key = self.get_name_and_sku(psu)
                if key in psu_dict:
                    psu_dict[key] += 1
                else:
                    psu_dict.update({key: 1})
            psu_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in psu_dict.items()])
            rows.append([_("Power Supplies"), psu_models])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class BladeUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, blade, centered=False):
        rows = [[_("Description"),_("Value")]]
        rows.append([_("Blade ID"), blade.id])
        rows.append([_("SKU"), blade.sku])
        rows.append([_("Model"), blade.name])
        rows.append([_("Serial Number"), blade.serial])
        rows.append([_("Firmware Package"), blade.firmware_package_version])

        memory_info = ""
        if blade.memory_arrays:
            memory_dict = {}
            for array in blade.memory_arrays:
                for unit in array.memory_units:
                    if not unit.capacity:
                        continue
                    if unit.sku:
                        if unit.clock:
                            key = str(int(unit.capacity / 1024)) + 'GB ' + str(unit.type) + ' ' + str(
                                unit.clock) + 'MHz (' + unit.sku + ')'
                        else:
                            key = str(int(unit.capacity / 1024)) + 'GB ' + str(unit.type) + ' (' + unit.sku + ')'
                    else:
                        if unit.clock:
                            key = str(int(unit.capacity / 1024)) + 'GB ' + str(unit.type) + ' ' + str(
                                unit.clock) + 'MHz'
                        else:
                            key = str(int(unit.capacity / 1024)) + 'GB' + str(unit.type)
                    if key in memory_dict:
                        memory_dict[key] += 1
                    else:
                        memory_dict.update({key: 1})
            memory_info = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in memory_dict.items()])
            rows.append([_("Memory"), str(blade.memory_total_marketing) + "\n" + memory_info])

        if blade.adaptors:
            adaptor_dict = {}
            for adaptor in blade.adaptors:
                key = self.get_name_and_sku(adaptor)
                if key in adaptor_dict:
                    adaptor_dict[key] += 1
                else:
                    adaptor_dict.update({key: 1})
            adaptor_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in adaptor_dict.items()])
            rows.append([_("Adaptors"), adaptor_models])

        cpu_model = ""
        cores = 0
        if blade.cpus:
            cpu_dict = {}

            for cpu in blade.cpus:
                key = self.get_name_and_sku(cpu)
                if key in cpu_dict:
                    cpu_dict[key] += 1
                else:
                    cpu_dict.update({key: 1})
                if cpu.cores:
                    cores += int(cpu.cores)

            cpu_model = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in cpu_dict.items()])

        rows.append([_("CPUs"), cpu_model])
        rows.append([_("Cores"), cores])

        gpu_models = ""
        if blade.gpus:
            gpu_dict = {}
            for gpu in blade.gpus:
                key = self.get_name_and_sku(gpu)
                if key in gpu_dict:
                    gpu_dict[key] += 1
                else:
                    gpu_dict.update({key: 1})
            gpu_models = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in gpu_dict.items()])
            rows.append([_("GPUs"), gpu_models])

        storage = ""
        drives = ""
        if blade.storage_controllers:
            storage_dict = {}
            drives_dict = {}
            for controller in blade.storage_controllers:
                key = self.get_name_and_sku(controller)
                if key in storage_dict:
                    storage_dict[key] += 1
                else:
                    storage_dict.update({key: 1})

                for drive in controller.disks:
                    key_drive = self.get_name_and_sku(drive)
                    if key_drive in drives_dict:
                        drives_dict[key_drive] += 1
                    else:
                        drives_dict.update({key_drive: 1})

            for drive in blade.nvme_drives:
                key_drive = self.get_name_and_sku(drive)
                if key_drive in drives_dict:
                    drives_dict[key_drive] += 1
                else:
                    drives_dict.update({key_drive: 1})

            storage = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in storage_dict.items()])
            drives = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in drives_dict.items()])
        rows.append([_("Storage Controllers"), storage])
        rows.append([_("Drives"), drives])

        flash_cards = ""
        if blade.storage_flexflash_controllers:
            flash_dict = {}
            for controller in blade.storage_flexflash_controllers:
                for card in controller.flexflash_cards:
                    if card.capacity_marketing in flash_dict:
                        flash_dict[card.capacity_marketing] += 1
                    else:
                        flash_dict.update({card.capacity_marketing: 1})
            flash_cards = "\n".join([(str(i[1]) + "x " + str(i[0])) for i in flash_dict.items()])
        rows.append([_("FlexFlash SD Cards"), flash_cards])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class VlanUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, vlans, centered=False):
        rows = [[_("ID"), _("VLAN Name"), _("Fabric"), _("Multicast Policy"), _("Sharing Type")]]

        vlans.sort(key=lambda x: int(x.id), reverse=False)
        for vlan in vlans:
            rows.append([vlan.id, vlan.name, vlan.fabric, vlan.multicast_policy_name, vlan.sharing_type])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class VlanGroupUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, vlan_groups, centered=False):
        rows = [[_("Name"),_("VLANs"), _("Native VLAN"), _("LAN Uplink Ports"), _("LAN Port Channel"),
                 _("Org Permissions")]]

        for vlan_group in vlan_groups:
            vlans = str(vlan_group.vlans).replace("[", "").replace("'", "").replace("]", "")
            lan_up = str(vlan_group.lan_uplink_ports).replace("[", "").replace("'", "").replace("]", "")
            lan_pc = str(vlan_group.lan_port_channels).replace("[", "").replace("'", "").replace("]", "")
            org_perm = str(vlan_group.org_permissions).replace("[", "").replace("'", "").replace("]", "")
            rows.append([vlan_group.name, vlans, vlan_group.native_vlan, lan_up, lan_pc, org_perm])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class LanPortChannelUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, lan_port_channels, centered=False):
        rows = [[_("ID"), _("Fabric"), _("Interfaces"), _("Name"), _("Admin Speed"), _("Flow Control Policy"),
                 _("LACP Policy")]]

        # TODO : Put in order by the PC ID (now : 119, 120, 13, 14, ...)
        for lan_port_channel in lan_port_channels:

            interfaces = ""
            for interface in lan_port_channel.interfaces:
                if "aggr_id" in interface.keys() and interface["aggr_id"]:
                    interfaces += interface["slot_id"] + "/" + interface["port_id"] + "/" + interface["aggr_id"] + ", "
                else:
                    interfaces += interface["slot_id"] + "/" + interface["port_id"] + ", "
            if interfaces:
                interfaces = interfaces[:-2]  # remove the last ", " at the end of the string

            rows.append([lan_port_channel.pc_id, lan_port_channel.fabric,
                         interfaces, lan_port_channel.name, lan_port_channel.admin_speed,
                         lan_port_channel.flow_control_policy, lan_port_channel.lacp_policy])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class SanPortChannelUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, san_port_channels, centered=False):
        rows = [[_("ID"), _("Fabric"), _("Interfaces"), _("Name"), _("Admin Speed"), _("VSAN Fabric"),
                 _("VSAN")]]

        # TODO : Put in order by the PC ID (now : 119, 120, 13, 14, ...)
        for san_port_channel in san_port_channels:

            interfaces = ""
            for interface in san_port_channel.interfaces:
                interfaces += interface["slot_id"] + "/" + interface["port_id"] + ", "
            if interfaces:
                interfaces = interfaces[:-2]  # remove the last ", " at the end of the string

            rows.append([san_port_channel.pc_id, san_port_channel.fabric,
                         interfaces, san_port_channel.name, san_port_channel.admin_speed,
                         san_port_channel.vsan_fabric, san_port_channel.vsan])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class FcoePortChannelUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, fcoe_port_channels, centered=False):
        rows = [[_("ID"), _("Fabric"), _("Interfaces"), _("Name"), _("LACP Policy")]]

        # TODO : Put in order by the PC ID (now : 119, 120, 13, 14, ...)
        for fcoe_port_channel in fcoe_port_channels:

            interfaces = ""
            for interface in fcoe_port_channel.interfaces:
                if "aggr_id" in interface.keys() and interface["aggr_id"]:
                    interfaces += interface["slot_id"] + "/" + interface["port_id"] + "/" + interface["aggr_id"] + ", "
                else:
                    interfaces += interface["slot_id"] + "/" + interface["port_id"] + ", "
            if interfaces:
                interfaces = interfaces[:-2]  # remove the last ", " at the end of the string

            rows.append([fcoe_port_channel.pc_id, fcoe_port_channel.fabric,
                         interfaces, fcoe_port_channel.name, fcoe_port_channel.lacp_policy])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class QosSystemClassUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, qos_system_class, centered=False):
        rows = [[_("Priority"), _("State"), _("CoS"), _("Packet Drop"), _("Weight"), _("MTU"),
                 _("Multicast Optimized")]]

        priority_order = ["platinum", "gold", "silver", "bronze", "fc", "best-effort"]
        order = {key: i for i, key in enumerate(priority_order)}

        for priority in sorted(qos_system_class, key=lambda x: order[x.priority], reverse=False):
            priority_name = priority.priority.replace("-"," ").title() if priority.priority != "fc" else "Fibre Channel"
            rows.append([priority_name, priority.state, priority.cos,
                         priority.packet_drop, priority.weight, priority.mtu, priority.multicast_optimized])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)

class VsanUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, vsans, centered=False):
        rows = [[_("ID"), _("FCoE VLAN ID"),_("VSAN Name"), _("Fabric"), _("Zoning")]]

        vsans.sort(key=lambda x: int(x.id), reverse=False)
        for vsan in vsans:
            rows.append([vsan.id, vsan.fcoe_vlan_id, vsan.name, vsan.fabric, vsan.zoning])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class IomSectionInfoUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, iom, centered=False):
        rows = [[_("ID"),_("SKU"), _("Model"), _("Serial Number"), _("Firmware"),
                 _("Fabric ports used")]]

        for iom_unit in iom:
            rows.append([iom_unit.id, iom_unit.sku, iom_unit.name, iom_unit.serial, iom_unit.firmware_package_version,
                         len(iom_unit.ports)])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class PsuSectionInfoUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, psu, centered=False):
        rows = [[_("ID"),_("SKU"), _("Model"), _("Serial Number")]]

        for power_supply in psu:
            if hasattr(power_supply, 'name'):  # IMC doesn't have .name
                name = power_supply.name
            else:
                name = power_supply.model

            rows.append([power_supply.id, power_supply.sku, name, power_supply.serial])

        GenericReportTable.__init__(self,order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class LicensesSectionInfoUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, fi, centered=False):
        rows = [[_("SKU"), _("Total"), _("Default"), _("Used"), _("Available"), _("Status"),
                 _("Grace Period Used (days)")]]

        for lic in fi.licenses:
            rows.append([lic["sku"], lic["quantity"], lic["quantity_default"], lic["quantity_used"],
                         lic["quantity_available"], lic["status"], lic["grace_period_used_days"]])

        GenericReportTable.__init__(self,order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class GemSectionInfoUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, gem, centered=False):
        rows = [[_("ID"),_("SKU"), _("Model"), _("Serial Number")]]

        for expansion_module in gem:
            rows.append([expansion_module.id, expansion_module.sku, expansion_module.name, expansion_module.serial])

        GenericReportTable.__init__(self,order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class DiskSectionInfoUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, disks, centered=False):
        rows = [[_("ID"),_("SKU"), _("Drive Type"), _("Connection Protocol"), _("Size"), _("Block Size"), _("RPM")]]

        for disk in disks:
            connection_protocol = None
            if hasattr(disk, "link_speed"):
                if disk.link_speed not in ["unknown", "NA", None]:
                    connection_protocol = disk.connection_protocol + ' (' + str(disk.link_speed) + 'Gbps)'
                else:
                    connection_protocol = disk.connection_protocol
            rpm = None
            if hasattr(disk, "rotational_speed_marketing"):
                rpm = disk.rotational_speed_marketing if disk.rotational_speed_marketing != 0 else None
            block_size = None
            if hasattr(disk, "block_size"):
                block_size = disk.block_size
            rows.append([disk.id, disk.sku, disk.drive_type, connection_protocol, disk.size_marketing, block_size, rpm])

        GenericReportTable.__init__(self,order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class BladesSectionInfoUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, blades, centered=False):
        rows = [[ _("ID"), _("Model"), _("Serial Number"), _("RAM"), _("CPUs"), _("Cores"),
                 _("Adapters"), _("GPUs"), _("Drives"), _("SD Cards")]]

        for blade in blades:
            cores = ""
            if blade.cpus:
                if blade.cpus[0].cores:
                    cores = int(blade.cpus[0].cores)
                    if len(blade.cpus) > 1:
                        cores = int(blade.cpus[0].cores) * len(blade.cpus)

            adaptor_sum = 0
            adaptor_models = ""  # If all drives have the same capacity, we write it down
            if blade.adaptors:
                adaptor_sum = len(blade.adaptors)
                if type(adaptor_models) == str:
                    adaptor_models = blade.adaptors[0].short_name
                if adaptor_models:
                    for adaptor in blade.adaptors:
                        if adaptor.short_name != adaptor_models:
                            adaptor_models = None
            if adaptor_models and adaptor_sum:
                adaptor_sum = str(adaptor_sum) + "x " + adaptor_models

            drives = 0
            drives_capacity = ""  # If all drives have the same capacity, we write it down
            for storage_controller in blade.storage_controllers:
                if storage_controller.disks:
                    drives += len(storage_controller.disks)
                    if type(drives_capacity) == str:
                        drives_capacity = storage_controller.disks[0].size_marketing
                    if drives_capacity:
                        for disk in storage_controller.disks:
                            if disk.size_marketing != drives_capacity:
                                drives_capacity = None
            if blade.nvme_drives:
                drives += len(blade.nvme_drives)
                if type(drives_capacity) == str:
                    drives_capacity = blade.nvme_drives[0].size_marketing
                if drives_capacity:
                    for disk in blade.nvme_drives:
                        if disk.size_marketing != drives_capacity:
                            drives_capacity = None
            if drives_capacity and drives:
                drives = str(drives) + "x " + drives_capacity

            sd_cards = 0
            sd_cards_capacity = ""   # If all drives have the same capacity, we write it down
            for storage_flexflash_controller in blade.storage_flexflash_controllers:
                if storage_flexflash_controller.flexflash_cards:
                    sd_cards += len(storage_flexflash_controller.flexflash_cards)
                    if type(sd_cards_capacity) == str:
                        sd_cards_capacity = storage_flexflash_controller.flexflash_cards[0].capacity_marketing
                    if sd_cards_capacity:
                        for sd_card in storage_flexflash_controller.flexflash_cards:
                            if sd_card.capacity_marketing != sd_cards_capacity:
                                sd_cards_capacity = None
                if sd_cards_capacity and sd_cards:
                    sd_cards = str(sd_cards) + "x " + sd_cards_capacity

            if parent.__class__.__name__ == "BladesUcsReportSection":
                id = blade.id
            else:
                id = blade.slot_id

            if blade.cpus:
                if blade.cpus[0].model_short_name:
                    rows.append([id, blade.short_name, blade.serial, blade.memory_total_marketing,
                                 str(len(blade.cpus)) + "x " + blade.cpus[0].model_short_name, cores, adaptor_sum,
                                 len(blade.gpus), drives, sd_cards])
                else:
                    rows.append([id, blade.short_name, blade.serial, blade.memory_total_marketing,
                                 str(len(blade.cpus)), cores, adaptor_sum, len(blade.gpus), drives, sd_cards])

        GenericReportTable.__init__(self,order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows, font_size=9)


class RackUnitsSectionInfoUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, rack_units, centered=False):
        rows = [[ _("ID"), _("Model"), _("Serial Number"), _("RAM"), _("CPUs"), _("Cores"),
                 _("Adapters"), _("GPUs"), _("Drives"), _("SD Cards")]]

        for rack in rack_units:
            cores = ""
            if rack.cpus:
                if rack.cpus[0].cores:
                    cores = int(rack.cpus[0].cores)
                    if len(rack.cpus) > 1:
                        cores = int(rack.cpus[0].cores) * len(rack.cpus)

            adaptor_sum = 0
            adaptor_models = ""  # If all drives have the same capacity, we write it down
            if rack.adaptors:
                adaptor_sum = len(rack.adaptors)
                if type(adaptor_models) == str:
                    adaptor_models = rack.adaptors[0].short_name
                if adaptor_models:
                    for adaptor in rack.adaptors:
                        if adaptor.short_name != adaptor_models:
                            adaptor_models = None
            if adaptor_models and adaptor_sum:
                adaptor_sum = str(adaptor_sum) + "x " + adaptor_models

            drives = 0
            drives_capacity = ""  # If all drives have the same capacity, we write it down
            for storage_controller in rack.storage_controllers:
                if storage_controller.disks:
                    drives += len(storage_controller.disks)
                    if type(drives_capacity) == str:
                        drives_capacity = storage_controller.disks[0].size_marketing
                    if drives_capacity:
                        for disk in storage_controller.disks:
                            if disk.size_marketing != drives_capacity:
                                drives_capacity = None
            if rack.nvme_drives:
                drives += len(rack.nvme_drives)
                if type(drives_capacity) == str:
                    drives_capacity = rack.nvme_drives[0].size_marketing
                if drives_capacity:
                    for disk in rack.nvme_drives:
                        if disk.size_marketing != drives_capacity:
                            drives_capacity = None
            if drives_capacity and drives:
                drives = str(drives) + "x " + drives_capacity

            sd_cards = 0
            sd_cards_capacity = ""   # If all drives have the same capacity, we write it down
            for storage_flexflash_controller in rack.storage_flexflash_controllers:
                if storage_flexflash_controller.flexflash_cards:
                    sd_cards += len(storage_flexflash_controller.flexflash_cards)
                    if type(sd_cards_capacity) == str:
                        sd_cards_capacity = storage_flexflash_controller.flexflash_cards[0].capacity_marketing
                    if sd_cards_capacity:
                        for sd_card in storage_flexflash_controller.flexflash_cards:
                            if sd_card.capacity_marketing != sd_cards_capacity:
                                sd_cards_capacity = None
                if sd_cards_capacity and sd_cards:
                    sd_cards = str(sd_cards) + "x " + sd_cards_capacity

            if rack.cpus:
                if rack.cpus[0].model_short_name:
                    rows.append([rack.id, rack.short_name, rack.serial, rack.memory_total_marketing,
                                 str(len(rack.cpus)) + "x " + rack.cpus[0].model_short_name, cores, adaptor_sum,
                                 len(rack.gpus), drives, sd_cards])
                else:
                    rows.append([rack.id, rack.short_name, rack.serial, rack.memory_total_marketing,
                                 str(len(rack.cpus)), cores, adaptor_sum, len(rack.gpus), drives, sd_cards])

        GenericReportTable.__init__(self,order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows, font_size=9)


class UcsReportRackConnectivityTable(GenericReportTable):
    def __init__(self, order_id, parent, fis, fexs, rack, centered=False, autofit=True):
        rows = [
            [_("Fabric"), _("Fabric Port"), _("FEX ID"), _("FEX Fabric Port"), _("FEX Host Port"),
             _("Adaptor ID"), _("Adaptor Port"), _("Transceiver SKU/Type"),
             _("Transceiver S/N"), _("Transceiver Length")]]

        for fi in fis:
            for port in fi.ports:
                if hasattr(port, 'peer'):
                    if port.peer:
                        fi_id = fi.id
                        if port.aggr_port_id:
                            fi_port = port.slot_id + "/" + port.aggr_port_id + "/" + port.port_id
                        else:
                            fi_port = port.slot_id + "/" + port.port_id

                        if "rack" in port.peer.keys():
                            if str(port.peer['rack']) == rack.id:
                                adaptor = port.peer['slot']
                                adaptor_port = port.peer['port']
                                if len(port.transceivers) == 1:
                                    if port.transceivers[0].sku:
                                        transceiver_type = port.transceivers[0].sku
                                    else:
                                        transceiver_type = port.transceivers[0].type
                                    transceiver_sn = port.transceivers[0].serial
                                    transceiver_length = port.transceivers[0].length
                                else:
                                    transceiver_type = ""
                                    transceiver_sn = ""
                                    transceiver_length = ""
                                rows.append(
                                    [fi_id, fi_port, "", "", "",
                                     adaptor, adaptor_port, transceiver_type,
                                     transceiver_sn,
                                     transceiver_length])

                        if "fex" in port.peer.keys():
                            for fex in fexs:
                                # We need to find the FEX and the associated peer port
                                if fex.id == str(port.peer['fex']):
                                    fex_fabric_port = str(port.peer['slot']) + "/" + str(port.peer['port'])
                                    for fex_host_port in fex.host_ports:
                                        if fex_host_port.peer:
                                            if "rack" in fex_host_port.peer.keys():
                                                if str(fex_host_port.peer['rack']) == rack.id:
                                                    fex_id = fex.id
                                                    if fex_host_port.aggr_port_id:
                                                        fex_port = fex_host_port.slot_id + "/" + \
                                                                   fex_host_port.aggr_port_id + "/" + \
                                                                   fex_host_port.port_id
                                                    else:
                                                        fex_port = fex_host_port.slot_id + "/" + fex_host_port.port_id
                                                    adaptor = fex_host_port.peer['slot']
                                                    adaptor_port = fex_host_port.peer['port']
                                                    if len(fex_host_port.transceivers) == 1:
                                                        if fex_host_port.transceivers[0].sku:
                                                            transceiver_type = fex_host_port.transceivers[0].sku
                                                        else:
                                                            transceiver_type = fex_host_port.transceivers[0].type
                                                        transceiver_sn = fex_host_port.transceivers[0].serial
                                                        transceiver_length = fex_host_port.transceivers[0].length
                                                    else:
                                                        transceiver_type = ""
                                                        transceiver_sn = ""
                                                        transceiver_length = ""
                                                    rows.append(
                                                        [fi_id, fi_port, fex_id, fex_fabric_port, fex_port,
                                                         adaptor, adaptor_port, transceiver_type,
                                                         transceiver_sn,
                                                         transceiver_length])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                        column_number=len(rows[1]), centered=centered, cells_list=rows, autofit=autofit)


class UcsReportChassisConnectivityTable(GenericReportTable):
    def __init__(self, order_id, parent, fis, fexs, chassis, centered=False, autofit=True):
        ext_port_type = "SIOC" if chassis.model in ["UCSC-C3X60", "UCSS-S3260"] else "IOM"
        rows = [[_("Fabric"), _("Fabric Port"), _("FEX ID"), _("FEX Fabric Port"), _("FEX Host Port"),_(ext_port_type),
                 _(ext_port_type + " Port"), _("Transceiver SKU/Type"), _("Transceiver S/N"),
                 _("Transceiver Length")]]

        for fi in fis:
            for port in fi.ports:
                if hasattr(port, 'peer'):
                    if port.peer:
                        fi_id = fi.id
                        if port.aggr_port_id:
                            fi_port = port.slot_id + "/" + port.aggr_port_id + "/" + port.port_id
                        else:
                            fi_port = port.slot_id + "/" + port.port_id

                        if "chassis" in port.peer.keys():
                            if str(port.peer['chassis']) == chassis.id:
                                iom = port.peer['slot']
                                iom_port = port.peer['port']
                                if len(port.transceivers) == 1:
                                    if port.transceivers[0].sku:
                                        transceiver_type = port.transceivers[0].sku
                                    else:
                                        transceiver_type = port.transceivers[0].type
                                    transceiver_sn = port.transceivers[0].serial
                                    transceiver_length = port.transceivers[0].length
                                else:
                                    transceiver_type = ""
                                    transceiver_sn = ""
                                    transceiver_length = ""
                                rows.append([fi_id, fi_port, "", "", "", iom, iom_port,
                                             transceiver_type, transceiver_sn,
                                             transceiver_length])

                        if "fex" in port.peer.keys():
                            for fex in fexs:
                                # We need to find the FEX and the associated peer port
                                if fex.id == str(port.peer['fex']):
                                    fex_fabric_port = str(port.peer['slot']) + "/" + str(port.peer['port'])
                                    for fex_host_port in fex.host_ports:
                                        if fex_host_port.peer:
                                            if "chassis" in fex_host_port.peer.keys():
                                                if str(fex_host_port.peer['rack']) == chassis.id:
                                                    fex_id = fex.id
                                                    if fex_host_port.aggr_port_id:
                                                        fex_port = fex_host_port.slot_id + "/" + \
                                                                   fex_host_port.aggr_port_id + "/" + \
                                                                   fex_host_port.port_id
                                                    else:
                                                        fex_port = fex_host_port.slot_id + "/" + fex_host_port.port_id
                                                    adaptor = fex_host_port.peer['slot']
                                                    adaptor_port = fex_host_port.peer['port']
                                                    if len(fex_host_port.transceivers) == 1:
                                                        if fex_host_port.transceivers[0].sku:
                                                            transceiver_type = fex_host_port.transceivers[0].sku
                                                        else:
                                                            transceiver_type = fex_host_port.transceivers[0].type
                                                        transceiver_sn = fex_host_port.transceivers[0].serial
                                                        transceiver_length = fex_host_port.transceivers[0].length
                                                    else:
                                                        transceiver_type = ""
                                                        transceiver_sn = ""
                                                        transceiver_length = ""
                                                    rows.append(
                                                        [fi_id, fi_port, fex_id, fex_fabric_port, fex_port,
                                                         adaptor, adaptor_port, transceiver_type,
                                                         transceiver_sn,
                                                         transceiver_length])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows, autofit=autofit)


class UcsReportLanNeighborsConnectivityTable(GenericReportTable):
    def __init__(self, order_id, parent, fis, centered=False, autofit=True):
        rows = [[_("Fabric"), _("Fabric Port"), _("Remote Name"), _("Remote Interface"),
                 _("Remote Model"), _("Remote S/N"), _("Transceiver SKU/Type"), _("Transceiver S/N"),
                 _("Transceiver Length")]]

        for fi in fis:
            for port in fi.ports:
                if port.type == "lan":
                    if hasattr(port, 'neighbor_entries'):
                        if port.neighbor_entries:
                            for neighbor_entry in port.neighbor_entries:
                                fi_id = fi.id
                                if port.aggr_port_id:
                                    fi_port = port.slot_id + "/" + port.aggr_port_id + "/" + port.port_id
                                else:
                                    fi_port = port.slot_id + "/" + port.port_id

                                remote_name = ""
                                if hasattr(neighbor_entry, "lldp_system_name"):
                                    if neighbor_entry.lldp_system_name:
                                        remote_name = neighbor_entry.lldp_system_name
                                if hasattr(neighbor_entry, "cdp_system_name") and not remote_name:
                                    if neighbor_entry.cdp_system_name:
                                        remote_name = neighbor_entry.cdp_system_name

                                remote_interface = ""
                                if hasattr(neighbor_entry, "cdp_remote_interface"):
                                    if neighbor_entry.cdp_remote_interface:
                                        remote_interface = neighbor_entry.cdp_remote_interface
                                if hasattr(neighbor_entry, "lldp_remote_interface") and not remote_interface:
                                    if neighbor_entry.lldp_remote_interface:
                                        remote_interface = neighbor_entry.lldp_remote_interface

                                remote_device_model = neighbor_entry.model
                                remote_sn = neighbor_entry.serial

                                if len(port.transceivers) == 1:
                                    if port.transceivers[0].sku:
                                        transceiver_type = port.transceivers[0].sku
                                    else:
                                        transceiver_type = port.transceivers[0].type
                                    transceiver_sn = port.transceivers[0].serial
                                    transceiver_length = port.transceivers[0].length
                                else:
                                    transceiver_type = ""
                                    transceiver_sn = ""
                                    transceiver_length = ""
                                rows.append([fi_id, fi_port, remote_name, remote_interface,
                                             remote_device_model, remote_sn, transceiver_type, transceiver_sn,
                                             transceiver_length])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows, autofit=autofit,
                                    font_size=8)


class UcsReportSanNeighborsConnectivityTable(GenericReportTable):
    def __init__(self, order_id, parent, fis, centered=False, autofit=True):
        rows = [[_("Fabric"), _("Fabric Port"), _("Remote Mgmt IP"), _("Remote WWNN"), _("Remote WWPN"),
                 _("Transceiver SKU/Type"), _("Transceiver S/N"), _("Transceiver Length")]]

        for fi in fis:
            for port in fi.ports:
                if port.type == "san":
                    if hasattr(port, 'neighbor_entries'):
                        if port.neighbor_entries:
                            for neighbor_entry in port.neighbor_entries:
                                fi_id = fi.id
                                if port.aggr_port_id:
                                    fi_port = port.slot_id + "/" + port.aggr_port_id + "/" + port.port_id
                                else:
                                    fi_port = port.slot_id + "/" + port.port_id

                                fabric_mgmt_addr = neighbor_entry.fabric_mgmt_addr \
                                    if neighbor_entry.fabric_mgmt_addr != "null" else None
                                fabric_nwwn = neighbor_entry.fabric_nwwn
                                fabric_pwwn = neighbor_entry.fabric_pwwn

                                if len(port.transceivers) == 1:
                                    if port.transceivers[0].sku:
                                        transceiver_type = port.transceivers[0].sku
                                    else:
                                        transceiver_type = port.transceivers[0].type
                                    transceiver_sn = port.transceivers[0].serial
                                    transceiver_length = port.transceivers[0].length
                                else:
                                    transceiver_type = ""
                                    transceiver_sn = ""
                                    transceiver_length = ""
                                rows.append([fi_id, fi_port, fabric_mgmt_addr, fabric_nwwn, fabric_pwwn,
                                             transceiver_type, transceiver_sn, transceiver_length])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows, autofit=autofit,
                                    font_size=8)


class ClusterInfoUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, config, device, centered=False):
        config_system = config.system[0]
        config_mng_int = config.management_interfaces[0]

        rows = [[_("Description"),_("Value")]]
        # Cluster info
        rows.append([_("System Name"), config_system.name])
        rows.append([_("Version"), config.device_version])
        rows.append([_("Cluster IP Address"),config_system.virtual_ip])
        rows.append([_("Netmask"), config_mng_int.netmask])
        rows.append([_("Gateway"), config_mng_int.gateway])
        if config_system.virtual_ipv6 and config_system.virtual_ipv6 not in ["::"]:
            rows.append([_("Cluster IP V6 Address"), config_system.virtual_ipv6])
        if config_mng_int.gateway_v6 and config_mng_int.gateway_v6 not in ["::"]:
            rows.append([_("Gateway V6"), config_mng_int.gateway_v6])
        # FI A (& B) info
        if config.management_interfaces:
            for interface in config.management_interfaces:
                if interface.fabric == "A":
                    if interface.ipv6 not in ["", "::"]:
                        rows.append([_("FI A - IPv6 Address"), interface.ipv6])
                    rows.append([_("FI A - IP Address"), interface.ip])

                if interface.fabric == "B":
                    if interface.ipv6 not in ["", "::"]:
                        rows.append([_("FI B - IPv6 Address"), interface.ipv6])
                    rows.append([_("FI B - IP Address"), interface.ip])
        # General cluster info
        if config_system.descr:
            rows.append([_("Description"), config_system.descr])
        if config_system.site:
            rows.append([_("Site"), config_system.site])
        if config_system.owner:
            rows.append([_("Owner"), config_system.owner])
        if config_system.domain_name:
            rows.append([_("Domain Name"), config_system.domain_name])
        # Timezone and DNS info
        if config.timezone_mgmt:
            if config.timezone_mgmt[0].zone:
                rows.append([_("Time Zone"), config.timezone_mgmt[0].zone])
            if config.timezone_mgmt[0].ntp:
                rows.append([_("NTP"),
                             str(config.timezone_mgmt[0].ntp).replace("'", "").replace("[", "").replace("]", "")])
        if config.dns:
            rows.append([_("DNS"), str(config.dns).replace("'", "").replace("[", "").replace("]", "")])
        if config.call_home:
            rows.append([_("Call Home"), config.call_home[0].admin_state])
        if config.ucs_central:
            rows.append([_("Link to UCS Central"), _("Registered to ") + config.ucs_central[0].ip_address])
        else:
            rows.append([_("Link to UCS Central"), "off"])
        rows.append([_("Intersight Claim Status"), config.intersight_status])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[0]), centered=centered, cells_list=rows)


class CommServicesInfoUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, config, device, centered=False):
        communication_services = config.communication_services[0]

        rows = [[_("Service"),_("State")]]
        # Cluster info
        rows.append([_("HTTP"), communication_services.http_service[0]["state"]])
        rows.append([_("Redirect to HTTPS"), communication_services.http_service[0]["redirect_to_https"]])
        rows.append([_("HTTPS"), communication_services.https_service[0]["state"]])
        rows.append([_("SNMP"), communication_services.snmp_service[0]["state"]])
        rows.append([_("CIMC"), communication_services.cimc_web_service])
        rows.append([_("SSH"), communication_services.ssh_service])
        rows.append([_("Telnet"), communication_services.telnet_service])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[0]), centered=centered, cells_list=rows, autofit=True)


class ServiceProfilesRacksUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, rack_units, centered=False):

        rows = [[_("Rack ID"), _("Service Profile"), _("Service Profile Org"), _("Service Profile Template"),
                 _("Service Profile Template Org")]]
        # Rack Service Profile info
        for rack in rack_units:
            rack_id = rack.id
            ls_name = rack.service_profile_name
            ls_org = rack.service_profile_org
            ls_template = None
            ls_template_org = None

            if rack.service_profile_template_name:
                ls_template = rack.service_profile_template_name
                ls_template_org = rack.service_profile_template_org
            rows.append([rack_id, ls_name, ls_org, ls_template, ls_template_org])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[0]), centered=centered, cells_list=rows)


class ServiceProfilesRackEnclosuresUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, rack_enclosures, centered=False):

        rows = [[_("Rack Enclosure ID"), _("Rack ID"), _("Service Profile"), _("Service Profile Org"),
                 _("Service Profile Template"), _("Service Profile Template Org")]]
        # Rack Enclosures Service Profile info
        for rack_enclosure in rack_enclosures:
            for server_node in rack_enclosure.server_nodes:
                rack_enclosure_id = rack_enclosure.id
                rack_id = server_node.id
                ls_name = server_node.service_profile_name
                ls_org = server_node.service_profile_org
                ls_template = None
                ls_template_org = None

                if server_node.service_profile_template_name:
                    ls_template = server_node.service_profile_template_name
                    ls_template_org = server_node.service_profile_template_org
                rows.append([rack_enclosure_id, rack_id, ls_name, ls_org, ls_template, ls_template_org])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[0]), centered=centered, cells_list=rows)


class ServiceProfilesChassisUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, chassis, centered=False):

        rows = [[_("Chassis ID"), _("Blade ID"), _("Service Profile"), _("Service Profile Org"),
                 _("Service Profile Template"), _("Service Profile Template Org")]]
        # Blade Service Profile info
        for chassis_unit in chassis:
            for blade in chassis_unit.blades:
                chassis_id = chassis_unit.id
                blade_id = blade.slot_id
                ls_name = blade.service_profile_name
                ls_org = blade.service_profile_org
                ls_template = None
                ls_template_org = None

                if blade.service_profile_template_name:
                    ls_template = blade.service_profile_template_name
                    ls_template_org = blade.service_profile_template_org
                rows.append([chassis_id, blade_id, ls_name, ls_org, ls_template, ls_template_org])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class ServiceProfileUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, service_profile, centered=False):
        rows = [[_("Description"),_("Value")]]
        rows.append([_("Name"), service_profile.name])
        rows.append([_("Type"), service_profile.type])
        rows.append([_("Organization"), service_profile._parent._dn])
        rows.append([_("BIOS Policy"), service_profile.bios_policy])
        rows.append([_("Boot Policy"), service_profile.boot_policy])
        rows.append([_("Maintenance Policy"), service_profile.maintenance_policy])
        rows.append([_("Local Disk Configuration Policy"), service_profile.local_disk_configuration_policy])
        rows.append([_("Dynamic vNIC Connection Policy"), service_profile.dynamic_vnic_connection_policy])
        rows.append([_("LAN Connectivity Policy"), service_profile.lan_connectivity_policy])
        rows.append([_("SAN Connectivity Policy"), service_profile.san_connectivity_policy])
        rows.append([_("Placement Policy"), service_profile.placement_policy])
        rows.append([_("vMedia Policy"), service_profile.vmedia_policy])
        rows.append([_("Serial Over LAN Policy"), service_profile.serial_over_lan_policy])
        rows.append([_("Threshold Policy"), service_profile.threshold_policy])
        rows.append([_("Power Control Policy"), service_profile.power_control_policy])
        rows.append([_("Scrub Policy"), service_profile.scrub_policy])
        rows.append([_("KVM Management Policy"), service_profile.kvm_management_policy])
        rows.append([_("Graphics Card Policy"), service_profile.graphics_card_policy])
        rows.append([_("Power Sync Policy"), service_profile.power_sync_policy])
        rows.append([_("Storage Profile"), service_profile.storage_profile])
        rows.append([_("IPMI Access Profile"), service_profile.ipmi_access_profile])
        rows.append([_("UUID Pool"), service_profile.uuid_pool])
        rows.append([_("WWNN Pool"), service_profile.wwnn_pool])
        rows.append([_("Server Pool"), service_profile.server_pool])
        rows.append([_("Host Firmware Package"), service_profile.host_firmware_package])
        if hasattr(service_profile, "children"):
            rows.append([_("Instantiated Service Profiles"), service_profile.children.replace(", ", "\n")])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class FiPortsUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, fi, centered=False, autofit=True):
        rows = [[_("Fabric"), _("Port"), _("Port Role"), _("Speed"),
                 _("Transceiver SKU/Type"), _("Transceiver S/N"), _("Transceiver Length")]]

        for port in fi.ports:
            fi_id = fi.id
            if port.aggr_port_id:
                fi_port = port.slot_id + "/" + port.aggr_port_id + "/" + port.port_id
            else:
                fi_port = port.slot_id + "/" + port.port_id

            fi_port_role = port.role if port.role != "unknown" else _("not configured")
            fi_port_speed = port.oper_speed if port.oper_speed != "indeterminate" else ""

            if len(port.transceivers) == 1:
                if port.transceivers[0].sku:
                    transceiver_type = port.transceivers[0].sku
                else:
                    transceiver_type = port.transceivers[0].type
                transceiver_sn = port.transceivers[0].serial
                transceiver_length = port.transceivers[0].length
            else:
                transceiver_type = ""
                transceiver_sn = ""
                transceiver_length = ""

            if fi_port_role != "not configured":
                rows.append([fi_id, fi_port, fi_port_role, fi_port_speed, transceiver_type, transceiver_sn,
                             transceiver_length])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows, autofit=autofit,
                                    font_size=8)


class IpPoolSectionInfoBlockUcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, blocks_ipv4, centered=False):
        rows = [[ _("From"), _("To"), _("Subnet"), _("Gateway"), _("Primary DNS"), _("Secondary DNS"), _("Size")]]

        for block_ipv4 in blocks_ipv4:
            size = int(IPv4Address(block_ipv4["to"])) - int(IPv4Address(block_ipv4["from"])) + 1
            rows.append([block_ipv4["from"], block_ipv4["to"], block_ipv4["netmask"],
                         block_ipv4["gateway"], block_ipv4["primary_dns"],
                         block_ipv4["secondary_dns"], size])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class IpPoolSectionInfoBlockv6UcsReportTable(GenericReportTable):
    def __init__(self, order_id, parent, blocks_ipv6, centered=False):
        rows = [[ _("From"), _("To"), _("Prefix"), _("Gateway"), _("Primary DNS"), _("Secondary DNS"), _("Size")]]

        for block_ipv6 in blocks_ipv6:
            size = int(IPv6Address(block_ipv6["to"])) - int(IPv6Address(block_ipv6["from"]))
            rows.append([block_ipv6["from"], block_ipv6["to"], block_ipv6["prefix"],
                         block_ipv6["gateway"], block_ipv6["primary_dns"],
                         block_ipv6["secondary_dns"], size])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)


class GenericSectionInfoBlockUcsReportTable(GenericReportTable):
    # Used for WWNN, WWPN, WWXN, UUID and MAC Blocks
    def __init__(self, order_id, parent, blocks, centered=False):
        rows = [[ _("From"), _("To"), _("Size")]]

        for block in blocks:
            size = int(block["to"].replace("-", "").replace(":", ""), 16) - int(block["from"].replace(
                "-", "").replace(":", ""), 16) + 1
            rows.append([block["from"], block["to"], size])

        GenericReportTable.__init__(self, order_id=order_id, parent=parent, row_number=len(rows),
                                    column_number=len(rows[1]), centered=centered, cells_list=rows)
